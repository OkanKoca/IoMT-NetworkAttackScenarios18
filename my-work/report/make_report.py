#!/usr/bin/env python3
"""
make_report.py — build the report .docx with python-docx.

Format follows the existing report series (page setup, title block, Heading 1
colour and size, bullet lists, Table Grid tables, centred figures with 8.5 pt
italic captions), so the documents read as one set.

Why this reads a Markdown source instead of holding the text inline, as the
earlier report scripts did: this document is long and every number in it is
checked against report_numbers.json, so it is corrected often. Text held inside
Python string literals cannot be proof-read comfortably, and a diff of it shows
re-wrapped lines rather than changed sentences. Keeping the prose in a Markdown
file leaves the source readable and the revision history meaningful, while this
script owns the formatting -- which is the part that has to match the series.

The Markdown accepted here is deliberately a small subset, and unknown syntax is
reported rather than silently dropped: a converter that quietly skips a block it
does not understand produces a document that looks complete and is not.

  ##/###/####     Heading 1 / 2 / 3
  paragraphs      **bold**, *italic*, `code`
  - item          bullet list (one level)
  1. item         numbered list
  | a | b |       table with a header row
  ```...```       code block
  > text          note block
  ---             horizontal rule (skipped; sections carry the separation)
  ![cap](path)    centred figure with an italic caption

Usage:
  python3 make_report.py                  # -> IoMT-Ag-Saldiri-Tespiti-Raporu.docx
  python3 make_report.py --out other.docx
  python3 make_report.py --force          # rebuild over a hand-edited .docx

The output is written from scratch every time, so anything added to the .docx by hand
-- a cover page, a table of contents, page numbers, a reviewer's comment -- is lost on
the next build. The final pass on this document is meant to be exactly those things,
so the build refuses to run once the .docx is newer than its Markdown source and says
what it found. --force overrides it, which is the right answer when the hand edits
have been folded back into the Markdown and the wrong one otherwise.
"""

import argparse
import re
import sys
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).resolve().parent
SOURCE = HERE / "IoMT-ag-saldiri-tespiti-raporu.md"
OUT = HERE / "IoMT-Ag-Saldiri-Tespiti-Raporu.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x40, 0x40, 0x40)
META = RGBColor(0x55, 0x55, 0x55)
CODE = RGBColor(0x30, 0x30, 0x30)


# --------------------------------------------------------------- document setup
def new_document():
    d = docx.Document()

    s = d.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.left_margin = s.right_margin = Inches(1.10)
    s.top_margin = s.bottom_margin = Inches(1.0)

    normal = d.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, colour, before in (("Heading 1", 14, ACCENT, 24),
                                       ("Heading 2", 11.5, DARK, 14),
                                       ("Heading 3", 10.5, DARK, 10)):
        st = d.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = colour
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(2)
    d.styles["Heading 3"].font.italic = True
    return d


def shade(paragraph, hex_fill):
    """Fill a paragraph's background. python-docx has no API for it, so the
    shading element is written into the paragraph properties directly."""
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    paragraph._p.get_or_add_pPr().append(el)


# --------------------------------------------------------------- inline runs
INLINE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)", re.S)


def add_runs(paragraph, text, base_italic=False):
    """Render **bold**, *italic* and `code` into runs of one paragraph."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = paragraph.add_run(piece[2:-2])
            r.font.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = paragraph.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = CODE
        elif piece.startswith("*") and piece.endswith("*"):
            r = paragraph.add_run(piece[1:-1])
            r.font.italic = True
        else:
            r = paragraph.add_run(piece)
        if base_italic:
            r.font.italic = True


# --------------------------------------------------------------- block writers
def title_block(d, title, meta_lines):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(19), True, ACCENT
    for line in meta_lines:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, line)
        for r in p.runs:
            r.font.size, r.font.color.rgb = Pt(9), META


def write_paragraph(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)


def write_list(d, items, numbered):
    for item in items:
        p = d.add_paragraph(style="List Number" if numbered else "List Bullet")
        add_runs(p, item)


def write_note(d, lines):
    """A '>' block: the short statement of what a section establishes."""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, " ".join(lines), base_italic=True)
    for r in p.runs:
        r.font.color.rgb = META
    shade(p, "F2F4F7")


def write_code(d, lines):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    r.font.color.rgb = CODE
    shade(p, "F4F4F4")


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def write_table(d, rows):
    header, body = split_row(rows[0]), [split_row(r) for r in rows[2:]]
    t = d.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        add_runs(cell.paragraphs[0], text)
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(9.5)
        shade(cell.paragraphs[0], "E8ECF1")
    for row in body:
        cells = t.add_row().cells
        for i, text in enumerate(row[:len(header)]):
            cells[i].text = ""
            add_runs(cells[i].paragraphs[0], text)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(9.5)
    d.add_paragraph()


def write_figure(d, caption, path, width=6.0):
    if not path.exists():
        raise FileNotFoundError(f"figure not found: {path}")
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    c = d.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.font.size, r.font.italic, r.font.color.rgb = Pt(8.5), True, META


# --------------------------------------------------------------- the parser
FIGURE = re.compile(r"^!\[(?P<cap>.*)\]\((?P<path>[^)]+)\)\s*$")
BULLET = re.compile(r"^[-*]\s+(?P<text>.*)$")
NUMBER = re.compile(r"^\d+\.\s+(?P<text>.*)$")


def build(source, out):
    lines = source.read_text(encoding="utf-8").split("\n")
    d = new_document()
    i, n = 0, len(lines)
    placed_title = False
    unknown = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # title block: the '# ' line plus the metadata lines up to the first rule
        if stripped.startswith("# ") and not placed_title:
            title = stripped[2:].strip()
            i += 1
            meta = []
            while i < n and lines[i].strip() != "---":
                if lines[i].strip():
                    meta.append(lines[i].strip())
                i += 1
            title_block(d, title, meta)
            placed_title = True
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            p = d.add_heading("", level=min(level - 1, 3))
            add_runs(p, text)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            block = []
            while i < n and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            write_code(d, block)
            continue

        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(lines[i])
                i += 1
            if len(rows) >= 2:
                write_table(d, rows)
            continue

        if stripped.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            write_note(d, [b for b in block if b])
            continue

        m = FIGURE.match(stripped)
        if m:
            write_figure(d, m["cap"], (source.parent / m["path"]).resolve())
            i += 1
            continue

        if BULLET.match(stripped) or NUMBER.match(stripped):
            numbered = bool(NUMBER.match(stripped))
            items = []
            while i < n:
                s = lines[i].strip()
                m = NUMBER.match(s) if numbered else BULLET.match(s)
                if m:
                    items.append(m["text"])
                    i += 1
                elif s and lines[i].startswith(("  ", "\t")) and items:
                    items[-1] += " " + s          # continuation of the last item
                    i += 1
                else:
                    break
            write_list(d, items, numbered)
            continue

        # otherwise: a paragraph, running until a blank line or a new block
        block = []
        while i < n and lines[i].strip() and not lines[i].strip().startswith(
                ("#", "|", ">", "```", "![")):
            s = lines[i].strip()
            if (BULLET.match(s) or NUMBER.match(s)) and block:
                break
            block.append(s)
            i += 1
        if block:
            write_paragraph(d, " ".join(block))
        else:
            unknown.append((i + 1, lines[i]))
            i += 1

    if unknown:
        print("unhandled lines (nothing was written for these):", file=sys.stderr)
        for num, text in unknown:
            print(f"  {num}: {text[:90]}", file=sys.stderr)

    d.save(out)
    return d


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--source", default=str(SOURCE))
    ap.add_argument("--out", default=str(OUT))
    ap.add_argument("--force", action="store_true",
                    help="rebuild even if the .docx looks hand-edited")
    args = ap.parse_args()

    src = Path(args.source)
    if not src.exists():
        sys.exit(f"source not found: {src}")
    out = Path(args.out)

    # "Newer than the source" cannot be the test: a successful build always leaves the
    # .docx newer than the Markdown, so it would refuse every second run. What marks a
    # hand edit is the file changing when this script did not write it -- so the build
    # records what it produced, and compares against that.
    stamp = out.with_suffix(".docx.built")
    if out.exists() and stamp.exists() and not args.force:
        if stamp.read_text().strip() != f"{out.stat().st_mtime_ns}":
            sys.exit(
                f"REFUSING TO REBUILD: {out.name} has changed since this script last\n"
                f"wrote it, so it holds edits that are not in {src.name} and a rebuild\n"
                f"would discard them. Fold them back into the Markdown, or pass --force.")

    build(src, out)
    stamp.write_text(f"{out.stat().st_mtime_ns}\n")
    print(f"{out}  ({out.stat().st_size / 1024:.0f} KB)")
    print("Add the page numbers, table of contents and cover page in Word or "
          "LibreOffice; those are the parts a builder should not guess at.")


if __name__ == "__main__":
    main()
